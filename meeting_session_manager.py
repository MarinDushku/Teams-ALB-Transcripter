#!/usr/bin/env python3
"""
Meeting Session Manager
Handles meeting metadata, timing, and comprehensive session management
"""

import json
import os
from datetime import datetime, timedelta
from pathlib import Path
import threading
import time

class MeetingSessionManager:
    def __init__(self):
        self.session_data = {
            'meeting_id': None,
            'meeting_title': 'Teams Meeting',
            'start_time': None,
            'end_time': None,
            'duration': None,
            'participants': {},
            'transcript': [],
            'statistics': {
                'total_words': 0,
                'total_speakers': 0,
                'most_active_speaker': None,
                'average_confidence': 0,
                'language_detected': 'Albanian'
            }
        }
        
        # Session state
        self.is_active = False
        self.auto_save_enabled = True
        self.save_interval = 30  # seconds
        
        # File paths
        self.sessions_dir = Path("sessions")
        self.sessions_dir.mkdir(exist_ok=True)
        
        # Auto-save thread
        self.save_thread = None
        self.stop_save_thread = False
        
        print("📝 Meeting Session Manager initialized")
    
    def start_session(self, meeting_title=None, meeting_id=None):
        """Start a new meeting session"""
        if self.is_active:
            print("⚠ Session already active")
            return False
        
        # Initialize session
        self.session_data = {
            'meeting_id': meeting_id or self.generate_meeting_id(),
            'meeting_title': meeting_title or f"Teams Meeting {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            'start_time': datetime.now().isoformat(),
            'end_time': None,
            'duration': None,
            'participants': {},
            'transcript': [],
            'statistics': {
                'total_words': 0,
                'total_speakers': 0,
                'most_active_speaker': None,
                'average_confidence': 0,
                'language_detected': 'Albanian'
            }
        }
        
        self.is_active = True
        
        # Start auto-save thread
        if self.auto_save_enabled:
            self.start_auto_save()
        
        print(f"🚀 Session started: {self.session_data['meeting_title']}")
        return True
    
    def end_session(self):
        """End the current meeting session"""
        if not self.is_active:
            print("⚠ No active session")
            return False
        
        # Finalize session data
        self.session_data['end_time'] = datetime.now().isoformat()
        start_time = datetime.fromisoformat(self.session_data['start_time'])
        end_time = datetime.fromisoformat(self.session_data['end_time'])
        self.session_data['duration'] = str(end_time - start_time)
        
        # Calculate final statistics
        self.calculate_final_statistics()
        
        # Stop auto-save
        self.stop_auto_save()
        
        # Final save
        self.save_session()
        
        self.is_active = False
        print(f"⏹️ Session ended: {self.session_data['meeting_title']}")
        print(f"📊 Duration: {self.session_data['duration']}")
        print(f"👥 Participants: {self.session_data['statistics']['total_speakers']}")
        print(f"💬 Words: {self.session_data['statistics']['total_words']}")
        
        return True
    
    def generate_meeting_id(self):
        """Generate a unique meeting ID"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"meeting_{timestamp}"
    
    def add_participant(self, name, joined_at=None, method='unknown'):
        """Add a participant to the session"""
        if not self.is_active:
            return False
        
        joined_time = joined_at or datetime.now()
        
        self.session_data['participants'][name] = {
            'joined_at': joined_time.isoformat() if isinstance(joined_time, datetime) else joined_time,
            'left_at': None,
            'status': 'active',
            'speaking_time': 0,
            'word_count': 0,
            'detection_method': method,
            'transcript_entries': []
        }
        
        print(f"👤 Added participant: {name}")
        return True
    
    def remove_participant(self, name, left_at=None):
        """Mark a participant as left"""
        if not self.is_active or name not in self.session_data['participants']:
            return False
        
        left_time = left_at or datetime.now()
        
        participant = self.session_data['participants'][name]
        participant['left_at'] = left_time.isoformat() if isinstance(left_time, datetime) else left_time
        participant['status'] = 'left'
        
        # Calculate speaking duration
        if participant['joined_at']:
            joined = datetime.fromisoformat(participant['joined_at'])
            left = datetime.fromisoformat(participant['left_at'])
            duration = left - joined
            participant['session_duration'] = str(duration)
        
        print(f"👤 Participant left: {name}")
        return True
    
    def add_transcript_entry(self, speaker, text, timestamp=None, confidence=None):
        """Add a transcript entry"""
        if not self.is_active:
            return False
        
        entry_time = timestamp or datetime.now()
        
        entry = {
            'timestamp': entry_time.isoformat() if isinstance(entry_time, datetime) else entry_time,
            'speaker': speaker,
            'text': text,
            'word_count': len(text.split()),
            'confidence': confidence,
            'entry_id': len(self.session_data['transcript']) + 1
        }
        
        self.session_data['transcript'].append(entry)
        
        # Update participant stats
        if speaker in self.session_data['participants']:
            participant = self.session_data['participants'][speaker]
            participant['word_count'] += entry['word_count']
            participant['transcript_entries'].append(entry['entry_id'])
        
        # Update session statistics
        self.session_data['statistics']['total_words'] += entry['word_count']
        
        return True
    
    def update_participant_speaking_time(self, speaker, speaking_duration):
        """Update speaking time for a participant"""
        if speaker in self.session_data['participants']:
            self.session_data['participants'][speaker]['speaking_time'] += speaking_duration
    
    def calculate_final_statistics(self):
        """Calculate final session statistics"""
        stats = self.session_data['statistics']
        participants = self.session_data['participants']
        transcript = self.session_data['transcript']
        
        # Count unique speakers
        speakers = set(entry['speaker'] for entry in transcript if entry['speaker'] != 'System')
        stats['total_speakers'] = len(speakers)
        
        # Find most active speaker
        if participants:
            most_active = max(participants.items(), key=lambda x: x[1]['word_count'])
            stats['most_active_speaker'] = {
                'name': most_active[0],
                'word_count': most_active[1]['word_count'],
                'speaking_time': most_active[1]['speaking_time']
            }
        
        # Calculate average confidence
        confidences = [entry['confidence'] for entry in transcript if entry['confidence'] is not None]
        if confidences:
            stats['average_confidence'] = sum(confidences) / len(confidences)
        
        # Calculate participation rates
        total_words = stats['total_words']
        for name, participant in participants.items():
            if total_words > 0:
                participant['participation_rate'] = (participant['word_count'] / total_words) * 100
            else:
                participant['participation_rate'] = 0
    
    def get_session_summary(self):
        """Get a summary of the current session"""
        if not self.is_active:
            return None
        
        current_time = datetime.now()
        start_time = datetime.fromisoformat(self.session_data['start_time'])
        current_duration = current_time - start_time
        
        active_participants = [name for name, info in self.session_data['participants'].items() 
                             if info['status'] == 'active']
        
        return {
            'meeting_title': self.session_data['meeting_title'],
            'current_duration': str(current_duration),
            'active_participants': active_participants,
            'total_participants': len(self.session_data['participants']),
            'transcript_entries': len(self.session_data['transcript']),
            'total_words': self.session_data['statistics']['total_words']
        }
    
    def save_session(self, filename=None):
        """Save session to file"""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            meeting_name = self.session_data['meeting_title'].replace(' ', '_').replace(':', '')
            filename = f"session_{meeting_name}_{timestamp}.json"
        
        filepath = self.sessions_dir / filename
        
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(self.session_data, f, indent=2, ensure_ascii=False)
            
            print(f"💾 Session saved: {filepath}")
            return str(filepath)
        
        except Exception as e:
            print(f"❌ Save error: {e}")
            return None
    
    def load_session(self, filepath):
        """Load session from file"""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                self.session_data = json.load(f)
            
            print(f"📂 Session loaded: {filepath}")
            return True
        
        except Exception as e:
            print(f"❌ Load error: {e}")
            return False
    
    def export_transcript(self, format='txt', filename=None):
        """Export transcript in various formats"""
        if not self.session_data['transcript']:
            print("⚠ No transcript data to export")
            return None
        
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            meeting_name = self.session_data['meeting_title'].replace(' ', '_')
            filename = f"transcript_{meeting_name}_{timestamp}.{format}"
        
        filepath = self.sessions_dir / filename
        
        try:
            if format == 'txt':
                self.export_as_text(filepath)
            elif format == 'json':
                self.export_as_json(filepath)
            elif format == 'docx':
                self.export_as_docx(filepath)
            elif format == 'csv':
                self.export_as_csv(filepath)
            else:
                print(f"❌ Unsupported format: {format}")
                return None
            
            print(f"📄 Transcript exported: {filepath}")
            return str(filepath)
        
        except Exception as e:
            print(f"❌ Export error: {e}")
            return None
    
    def export_as_text(self, filepath):
        """Export as plain text"""
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"🎭 {self.session_data['meeting_title']}\n")
            f.write("=" * 60 + "\n\n")
            f.write(f"Start Time: {self.session_data['start_time']}\n")
            f.write(f"Duration: {self.session_data.get('duration', 'In progress')}\n")
            f.write(f"Participants: {len(self.session_data['participants'])}\n\n")
            
            f.write("TRANSCRIPT:\n")
            f.write("-" * 40 + "\n\n")
            
            for entry in self.session_data['transcript']:
                timestamp = entry['timestamp']
                if isinstance(timestamp, str) and 'T' in timestamp:
                    dt = datetime.fromisoformat(timestamp)
                    time_str = dt.strftime("%H:%M:%S")
                else:
                    time_str = str(timestamp)
                
                f.write(f"[{time_str}] {entry['speaker']}: {entry['text']}\n")
                if entry.get('confidence'):
                    f.write(f"  (Confidence: {entry['confidence']:.1%})\n")
                f.write("\n")
    
    def export_as_json(self, filepath):
        """Export as JSON"""
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(self.session_data, f, indent=2, ensure_ascii=False)
    
    def export_as_docx(self, filepath):
        """Export as Word document"""
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Title
            title = doc.add_heading(self.session_data['meeting_title'], 0)
            
            # Meeting info
            info = doc.add_paragraph()
            info.add_run("Start Time: ").bold = True
            info.add_run(f"{self.session_data['start_time']}\n")
            info.add_run("Duration: ").bold = True
            info.add_run(f"{self.session_data.get('duration', 'In progress')}\n")
            info.add_run("Participants: ").bold = True
            info.add_run(f"{len(self.session_data['participants'])}")
            
            # Participant list
            doc.add_heading("Participants", level=1)
            for name, info in self.session_data['participants'].items():
                p = doc.add_paragraph(f"• {name}")
                if info.get('participation_rate'):
                    p.add_run(f" ({info['participation_rate']:.1f}% participation)")
            
            # Transcript
            doc.add_heading("Transcript", level=1)
            for entry in self.session_data['transcript']:
                timestamp = entry['timestamp']
                if isinstance(timestamp, str) and 'T' in timestamp:
                    dt = datetime.fromisoformat(timestamp)
                    time_str = dt.strftime("%H:%M:%S")
                else:
                    time_str = str(timestamp)
                
                p = doc.add_paragraph()
                p.add_run(f"[{time_str}] ").bold = True
                p.add_run(f"{entry['speaker']}: ").bold = True
                p.add_run(entry['text'])
            
            doc.save(filepath)
        
        except ImportError:
            print("⚠ python-docx not available, exporting as text instead")
            self.export_as_text(filepath.with_suffix('.txt'))
    
    def export_as_csv(self, filepath):
        """Export as CSV"""
        import csv
        
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['Timestamp', 'Speaker', 'Text', 'Word Count', 'Confidence'])
            
            for entry in self.session_data['transcript']:
                writer.writerow([
                    entry['timestamp'],
                    entry['speaker'],
                    entry['text'],
                    entry['word_count'],
                    entry.get('confidence', '')
                ])
    
    def start_auto_save(self):
        """Start auto-save thread"""
        if self.save_thread and self.save_thread.is_alive():
            return
        
        self.stop_save_thread = False
        self.save_thread = threading.Thread(target=self.auto_save_loop, daemon=True)
        self.save_thread.start()
        print(f"💾 Auto-save enabled (every {self.save_interval}s)")
    
    def stop_auto_save(self):
        """Stop auto-save thread"""
        self.stop_save_thread = True
        if self.save_thread:
            self.save_thread.join(timeout=2)
    
    def auto_save_loop(self):
        """Auto-save loop"""
        while not self.stop_save_thread and self.is_active:
            time.sleep(self.save_interval)
            if self.is_active:
                self.save_session()
    
    def get_recent_sessions(self, limit=10):
        """Get list of recent sessions"""
        session_files = []
        
        for file_path in self.sessions_dir.glob("session_*.json"):
            try:
                stat = file_path.stat()
                session_files.append({
                    'filename': file_path.name,
                    'path': str(file_path),
                    'modified': datetime.fromtimestamp(stat.st_mtime),
                    'size': stat.st_size
                })
            except Exception:
                continue
        
        # Sort by modification time (newest first)
        session_files.sort(key=lambda x: x['modified'], reverse=True)
        
        return session_files[:limit]

# Test function
if __name__ == "__main__":
    manager = MeetingSessionManager()
    
    # Test session
    manager.start_session("Test Meeting")
    
    # Add some participants
    manager.add_participant("John Smith")
    manager.add_participant("Maria Doe")
    
    # Add some transcript entries
    manager.add_transcript_entry("John Smith", "Hello everyone, how are you today?", confidence=0.95)
    manager.add_transcript_entry("Maria Doe", "Mirëdita! I'm doing great, thanks!", confidence=0.92)
    
    # Show summary
    summary = manager.get_session_summary()
    print(f"📊 Session Summary: {summary}")
    
    # End session
    manager.end_session()
    
    # Export transcript
    manager.export_transcript('txt')