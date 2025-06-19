import tkinter as tk
from tkinter import ttk, scrolledtext, filedialog, messagebox
import threading
import json
from datetime import datetime

class LiveTranscriptUI:
    def __init__(self):
        self.window = tk.Tk()
        self.window.title("Albanian Teams Transcriber")
        self.window.geometry("1000x700")
        self.window.configure(bg='#f0f0f0')
        
        # Speaker colors
        self.speaker_colors = {
            "Speaker 1": "#0066cc",  # Blue
            "Speaker 2": "#009900",  # Green
            "Speaker 3": "#cc0000",  # Red
            "Speaker 4": "#ff6600",  # Orange
            "Speaker 5": "#9900cc",  # Purple
            "Speaker 6": "#cc6600",  # Brown
        }
        
        # Transcript data
        self.transcript_data = []
        self.is_recording = False
        
        self.setup_ui()
        
    def setup_ui(self):
        """Setup the user interface"""
        # Title
        title_label = tk.Label(
            self.window, 
            text="Albanian Teams Transcriber", 
            font=("Arial", 16, "bold"),
            bg='#f0f0f0'
        )
        title_label.pack(pady=10)
        
        # Control frame
        control_frame = tk.Frame(self.window, bg='#f0f0f0')
        control_frame.pack(pady=5)
        
        # Start/Stop button
        self.start_stop_btn = tk.Button(
            control_frame, 
            text="Start Auto-Transcription", 
            command=self.toggle_transcription,
            font=("Arial", 12),
            bg='#4CAF50',
            fg='white',
            width=20
        )
        self.start_stop_btn.pack(side=tk.LEFT, padx=5)
        
        # Clear button
        self.clear_btn = tk.Button(
            control_frame,
            text="Clear Transcript",
            command=self.clear_transcript,
            font=("Arial", 12),
            bg='#f44336',
            fg='white',
            width=15
        )
        self.clear_btn.pack(side=tk.LEFT, padx=5)
        
        # Export button
        self.export_btn = tk.Button(
            control_frame,
            text="Export Transcript",
            command=self.export_transcript,
            font=("Arial", 12),
            bg='#2196F3',
            fg='white',
            width=15
        )
        self.export_btn.pack(side=tk.LEFT, padx=5)
        
        # Status frame
        status_frame = tk.Frame(self.window, bg='#f0f0f0')
        status_frame.pack(pady=5)
        
        # Status label
        self.status_label = tk.Label(
            status_frame,
            text="Status: Ready",
            font=("Arial", 10),
            bg='#f0f0f0'
        )
        self.status_label.pack(side=tk.LEFT)
        
        # Teams detection indicator
        self.teams_indicator = tk.Label(
            status_frame,
            text="Teams: Not Detected",
            font=("Arial", 10),
            bg='#f0f0f0',
            fg='red'
        )
        self.teams_indicator.pack(side=tk.RIGHT)
        
        # Transcript display frame
        transcript_frame = tk.Frame(self.window)
        transcript_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # Transcript area
        self.transcript_area = scrolledtext.ScrolledText(
            transcript_frame, 
            wrap=tk.WORD, 
            width=100, 
            height=30,
            font=("Arial", 11),
            bg='white',
            relief=tk.SUNKEN,
            borderwidth=2
        )
        self.transcript_area.pack(fill=tk.BOTH, expand=True)
        
        # Configure text tags for different speakers
        for speaker, color in self.speaker_colors.items():
            self.transcript_area.tag_configure(
                speaker.lower().replace(' ', '_'), 
                foreground=color,
                font=("Arial", 11, "bold")
            )
        
        # Timestamp tag
        self.transcript_area.tag_configure(
            "timestamp", 
            foreground="#666666",
            font=("Arial", 9)
        )
        
        # Statistics frame
        stats_frame = tk.Frame(self.window, bg='#f0f0f0')
        stats_frame.pack(fill=tk.X, padx=10, pady=5)
        
        self.stats_label = tk.Label(
            stats_frame,
            text="Speakers: 0 | Transcribed: 0 minutes",
            font=("Arial", 9),
            bg='#f0f0f0'
        )
        self.stats_label.pack()
        
        # Callback functions
        self.start_callback = None
        self.stop_callback = None
        
    def set_callbacks(self, start_callback, stop_callback):
        """Set callback functions for start/stop actions"""
        self.start_callback = start_callback
        self.stop_callback = stop_callback
        
    def toggle_transcription(self):
        """Toggle transcription on/off"""
        if not self.is_recording:
            self.start_transcription()
        else:
            self.stop_transcription()
    
    def start_transcription(self):
        """Start transcription"""
        self.is_recording = True
        self.start_stop_btn.config(
            text="Stop Transcription",
            bg='#f44336'
        )
        self.status_label.config(text="Status: Recording...")
        
        if self.start_callback:
            self.start_callback()
    
    def stop_transcription(self):
        """Stop transcription"""
        self.is_recording = False
        self.start_stop_btn.config(
            text="Start Auto-Transcription",
            bg='#4CAF50'
        )
        self.status_label.config(text="Status: Stopped")
        
        if self.stop_callback:
            self.stop_callback()
    
    def update_transcript(self, speaker_id, text, timestamp=None):
        """Update transcript with new text"""
        if not timestamp:
            timestamp = datetime.now().strftime("%H:%M:%S")
        
        # Add to transcript data
        entry = {
            'timestamp': timestamp,
            'speaker': speaker_id or "Unknown",
            'text': text
        }
        self.transcript_data.append(entry)
        
        # Format display text
        formatted_text = f"[{timestamp}] {speaker_id or 'Unknown'}: {text}\n"
        
        # Insert with appropriate styling
        self.transcript_area.insert(tk.END, f"[{timestamp}] ", "timestamp")
        
        speaker_tag = (speaker_id or "unknown").lower().replace(' ', '_')
        if speaker_tag in [tag.replace(' ', '_') for tag in self.speaker_colors.keys()]:
            self.transcript_area.insert(tk.END, f"{speaker_id}: ", speaker_tag)
        else:
            self.transcript_area.insert(tk.END, f"{speaker_id}: ")
        
        self.transcript_area.insert(tk.END, f"{text}\n")
        self.transcript_area.see(tk.END)  # Auto-scroll to bottom
        
        # Update statistics
        self.update_statistics()
        
        # Auto-save
        self.auto_save_transcript()
    
    def update_statistics(self):
        """Update statistics display"""
        speakers = set(entry['speaker'] for entry in self.transcript_data)
        total_entries = len(self.transcript_data)
        
        # Estimate minutes (rough calculation)
        estimated_minutes = total_entries // 20 if total_entries > 0 else 0
        
        self.stats_label.config(
            text=f"Speakers: {len(speakers)} | Transcribed: ~{estimated_minutes} minutes | Total entries: {total_entries}"
        )
    
    def update_teams_status(self, is_detected, reason=""):
        """Update Teams detection status"""
        if is_detected:
            self.teams_indicator.config(
                text=f"Teams: Detected ({reason})",
                fg='green'
            )
        else:
            self.teams_indicator.config(
                text="Teams: Not Detected",
                fg='red'
            )
    
    def clear_transcript(self):
        """Clear the transcript"""
        if messagebox.askyesno("Clear Transcript", "Are you sure you want to clear the transcript?"):
            self.transcript_area.delete(1.0, tk.END)
            self.transcript_data = []
            self.update_statistics()
    
    def export_transcript(self):
        """Export transcript to file"""
        if not self.transcript_data:
            messagebox.showwarning("Export", "No transcript data to export!")
            return
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[
                ("Text files", "*.txt"),
                ("JSON files", "*.json"),
                ("Word documents", "*.docx"),
                ("All files", "*.*")
            ]
        )
        
        if file_path:
            try:
                if file_path.endswith('.json'):
                    self.export_json(file_path)
                elif file_path.endswith('.docx'):
                    self.export_docx(file_path)
                else:
                    self.export_text(file_path)
                
                messagebox.showinfo("Export", f"Transcript exported to {file_path}")
            except Exception as e:
                messagebox.showerror("Export Error", f"Failed to export: {e}")
    
    def export_text(self, file_path):
        """Export as plain text"""
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write("Albanian Teams Transcript\n")
            f.write("=" * 50 + "\n\n")
            
            for entry in self.transcript_data:
                f.write(f"[{entry['timestamp']}] {entry['speaker']}: {entry['text']}\n")
    
    def export_json(self, file_path):
        """Export as JSON"""
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump({
                'export_time': datetime.now().isoformat(),
                'transcript': self.transcript_data
            }, f, indent=2, ensure_ascii=False)
    
    def export_docx(self, file_path):
        """Export as Word document (requires python-docx)"""
        try:
            from docx import Document
            
            doc = Document()
            doc.add_heading('Albanian Teams Transcript', 0)
            
            for entry in self.transcript_data:
                p = doc.add_paragraph()
                p.add_run(f"[{entry['timestamp']}] ").bold = True
                p.add_run(f"{entry['speaker']}: ").italic = True
                p.add_run(entry['text'])
            
            doc.save(file_path)
        except ImportError:
            # Fallback to text export
            self.export_text(file_path.replace('.docx', '.txt'))
            messagebox.showwarning("Export", "python-docx not installed. Exported as text file instead.")
    
    def auto_save_transcript(self):
        """Auto-save transcript periodically"""
        if len(self.transcript_data) % 10 == 0:  # Save every 10 entries
            try:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"transcript_autosave_{timestamp}.json"
                self.export_json(filename)
            except:
                pass  # Silent fail for auto-save
    
    def run(self):
        """Start the UI main loop"""
        self.window.mainloop()
    
    def destroy(self):
        """Close the window"""
        self.window.destroy()